import time
from typing import Dict, List, Optional
import boto3
import openai
import os
import json
from datetime import datetime
from dotenv import load_dotenv
from openai import OpenAI
import tiktoken
import ast
from docx import Document
import numpy as np

from langchain.agents import initialize_agent, Tool
from langchain.prompts import PromptTemplate
from langchain.llms import OpenAI
from langchain.chains import LLMChain


import sys
current_directory = os.getcwd()
sys.path.append(current_directory)
app_dir = os.path.join(current_directory, "app")
sys.path.append(app_dir)


from app.utils.cost_tracker import QueryCostTracker


# Load .env variables
load_dotenv("app/.env")

# Initialize OpenAI API key (replace with your key)
openai.api_key = os.getenv("OPENAI_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_TRANSCRIPTION_PROJ_ID = "proj_Q8SXL9dUJqv3kBv3PUlotdt6"
OPENAI_ORG_NAME = "org-3m84stKBZ07lg0LvLzHGh5Yb"
# OPENAI_TRANSCRIPTION_PROJ_NAME = "Transcription-LLM"

openai_client = OpenAI(
    api_key=OPENAI_API_KEY,
    project=OPENAI_TRANSCRIPTION_PROJ_ID,
    organization=OPENAI_ORG_NAME,
)

# AWS Credentials from environment variables
AWS_ACCESS_KEY_ID = os.getenv("AWS_ACCESS_KEY_ID")
AWS_SECRET_ACCESS_KEY = os.getenv("AWS_SECRET_ACCESS_KEY")

# AWS default settings
DEFAULT_REGION_NAME = "us-west-2"
DEFAULT_BUCKET_NAME = "digital-resume-s3"

# AWS Transcribe constants
MEDIA_FORMAT = "wav"  # Options: "wav", "mp3", "mp4", "flac"
LANGUAGE_CODE = "en-ZA"  # Options: "en-US", "en-GB", "es-US", "fr-CA", etc.
DEFAULT_TRANSCRIBE_MAX_SPEAKERS = 3

# OpenAI constants
OPENAI_MODEL = "gpt-4"  # Options: "gpt-3.5-turbo", "text-davinci-003", "text-curie-001", "gpt-4"
OPENAI_MAX_TOKENS_SUMMARY = 300
OPENAI_MAX_TOKENS_EXTRACTION = 300
OPENAI_MAX_TOKENS_CLASSIFICATION = 500
OPENAI_TEMPERATURE = 0.5  # Options: Float between 0 (deterministic) and 1 (creative)


# Script Parameters
write_queries_to_log_file = True
DRY_RUN_SYSTEM = True

USD_TO_ZAR_CONVERSION = 19.10


DEFAULT_FP_NAME = "Morgan"

cost_tracker = QueryCostTracker(usd_to_zar_conversion_rate=USD_TO_ZAR_CONVERSION)


LOG_DIR = "logs"
LLM_LOG_FILE_NAME = f"{LOG_DIR}/llm_api_query_logs.txt"
FULL_APP_LOG_FILE_NAME = f"{LOG_DIR}/full_app_output.txt"

outputs_dir = "outputs"
transcript_out_dir = os.path.join(outputs_dir, "transcripts")


def log_final_results(final_results, stats, log_final_file: str = FULL_APP_LOG_FILE_NAME):
    """
    Logs the final results and statistics to a file with a timestamp.
    """

    with open(log_final_file, "a") as log_file:
        log_file.write(f"\n--- Log Entry: {datetime.now()} ---\n")
        log_file.write(f"Final Results:\n{final_results}\n")
        log_file.write(f"Statistics:\n{stats}\n")
        log_file.write("--- End of Entry ---\n")
        

class Boto3Client:
    """Manages Boto3 client connections."""
    def __init__(self, region: str = DEFAULT_REGION_NAME):
        self.transcribe_client : boto3.client = boto3.client(
            "transcribe",
            region_name=region,
            aws_access_key_id=AWS_ACCESS_KEY_ID,
            aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
        )
        self.s3_client = boto3.client(
            "s3",
            region_name=region,
            aws_access_key_id=AWS_ACCESS_KEY_ID,
            aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
        )

class S3Manager:
    """Handles S3 operations."""
    def __init__(self, boto3_client: Boto3Client, bucket_name: str = DEFAULT_BUCKET_NAME):
        self.s3_client = boto3_client.s3_client
        self.bucket_name = bucket_name

    def upload_file(self, local_file_path: str, s3_key: str) -> str:
        """Uploads a file to S3 and returns the S3 URI."""
        self.s3_client.upload_file(local_file_path, self.bucket_name, s3_key)
        return f"s3://{self.bucket_name}/{s3_key}"

class TranscriptionManager:
    """Manages transcription jobs."""
    def __init__(self, boto3_client: Boto3Client):
        self.transcribe_client = boto3_client.transcribe_client
        
    def get_client(self) -> boto3.client:
        return self.transcribe_client

    def wait_for_transcription_completion(self, job_name: str):
        """Waits until the transcription job completes."""
        while True:
            response = self.transcribe_client.get_transcription_job(
                TranscriptionJobName=job_name
            )
            status = response["TranscriptionJob"]["TranscriptionJobStatus"]
            if status in ["COMPLETED", "FAILED"]:
                break
            print("Waiting for transcription to complete...")
            time.sleep(10)

        if status == "FAILED":
            raise RuntimeError(f"Transcription job failed: {response}")
        return response

    def fetch_transcription_output(self, job_name: str) -> dict:
        """Fetches the transcription output JSON."""
        response = self.transcribe_client.get_transcription_job(
            TranscriptionJobName=job_name
        )
        uri = response["TranscriptionJob"]["Transcript"]["TranscriptFileUri"]
        file_name = "other/audio-files/output/" + uri.split("/")[-1]
        # file_name = "other/audio-files/output/" + uri.split("/")[-1]
        print("file_name", file_name)
        transcription_file_data = boto3.client("s3").get_object(
            Bucket=DEFAULT_BUCKET_NAME, 
            Key=file_name,
        )
        transcription_data = json.loads(transcription_file_data["Body"].read())['results']
        return transcription_data

def initialize_clients():
    """Initializes S3 and Transcription clients."""
    boto3_client = Boto3Client()
    s3_manager = S3Manager(boto3_client)
    transcription_manager = TranscriptionManager(boto3_client)
    return boto3_client, s3_manager, transcription_manager


def upload_audio_to_s3(s3_manager, local_file_path: str, directory_name: str = "other/audio-files") -> str:
    """
    Uploads an audio file to S3 and creates a structured folder hierarchy on S3.

    Parameters:
    - s3_manager: An S3 manager object with `s3_client` and `bucket_name`.
    - local_file_path (str): The path of the local file to upload.
    - directory_name (str): The base directory path inside the S3 bucket. Default is "other/audio-files".

    Returns:
    - str: The S3 URL of the uploaded file.
    """
    # Get the current date in dd-mm-yy format
    current_date = datetime.now().strftime("%d-%m-%y")
    job_folder = f"job_{current_date}"

    # Define S3 paths
    input_s3_key = f"{directory_name}/{job_folder}/input/{os.path.basename(local_file_path)}"
    output_s3_folder = f"{directory_name}/{job_folder}/output/"

    s3_client = s3_manager.s3_client

    # Check if the file already exists in S3
    try:
        s3_client.head_object(Bucket=s3_manager.bucket_name, Key=input_s3_key)
        print(f"File already exists in S3: s3://{s3_manager.bucket_name}/{input_s3_key}")
        return f"s3://{s3_manager.bucket_name}/{input_s3_key}"
    except s3_client.exceptions.ClientError as e:
        if e.response['Error']['Code'] == "404":
            print(f"File does not exist in S3. Proceeding with upload.")
            # pass
        else:
            print(f"Error checking file in S3: {e}")
            raise

    # Upload the file to the input directory on S3
    try:
        s3_client.upload_file(local_file_path, s3_manager.bucket_name, input_s3_key)
        print(f"Uploaded to S3: s3://{s3_manager.bucket_name}/{input_s3_key}")
    except Exception as e:
        print(f"Error uploading file to S3: {e}")
        raise

    # Create the output folder on S3 (S3 "folders" are just keys that end with '/')
    try:
        s3_client.put_object(Bucket=s3_manager.bucket_name, Key=output_s3_folder)
        print(f"Created output folder in S3: s3://{s3_manager.bucket_name}/{output_s3_folder}")
    except Exception as e:
        print(f"Error creating output folder in S3: {e}")
        raise

    return f"s3://{s3_manager.bucket_name}/{input_s3_key}"


VOCABULARY_DIR = "vocabulary"
TRANSCRIBE_VOCABULARY_FILE = f"{VOCABULARY_DIR}/FP-SA-Vocab-V1.csv"



def upload_vocabulary_files(transcribe_client, vocab_name: str, vocab_file: str):
    """
    Uploads a custom vocabulary file to AWS Transcribe.

    Parameters:
    - transcribe_client: The AWS Transcribe client.
    - vocab_name (str): The name of the vocabulary to create or update.
    - vocab_file (str): The local file path of the vocabulary CSV.

    Returns:
    - None
    """
    try:
        # Check if the vocabulary already exists
        response = transcribe_client.get_vocabulary(VocabularyName=vocab_name)
        if response['VocabularyState'] in ['READY', 'PENDING']:
            print(f"Vocabulary '{vocab_name}' already exists and is ready or pending. Skipping upload.")
            return
    except transcribe_client.exceptions.BadRequestException:
        print(f"Vocabulary '{vocab_name}' does not exist. Proceeding to create it.")


    raise("Error: I need to setup uploading file in .txt form to s3 first - Phrases does not work. Use VocabularyFileUri instead")
    # Upload vocabulary
    with open(vocab_file, 'r') as file:
        vocab_content = file.read()

    try:
        transcribe_client.create_vocabulary(
            VocabularyName=vocab_name,
            LanguageCode=LANGUAGE_CODE,
            Phrases=vocab_content.splitlines(),
        )
        print(f"Vocabulary '{vocab_name}' has been uploaded successfully.")
    except Exception as e:
        print(f"Error uploading vocabulary '{vocab_name}': {e}")
        raise


def start_aws_transcription_job(transcription_manager, input_s3_uri: str, output_s3_dir_prefix: str) -> str:
    """Starts a transcription job and returns the job name."""
    transcribe_client = transcription_manager.transcribe_client
    # Upload custom vocabularies
    upload_vocabulary_files(transcribe_client, "FP-SA-Vocab-V1", TRANSCRIBE_VOCABULARY_FILE)

    job_name = f"transcription_job_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

    try:
        transcribe_client.start_transcription_job(
            TranscriptionJobName=job_name,
            Media={"MediaFileUri": input_s3_uri},
            MediaFormat=MEDIA_FORMAT,
            LanguageCode=LANGUAGE_CODE,
            OutputBucketName=DEFAULT_BUCKET_NAME,
            OutputKey=output_s3_dir_prefix,
            Settings={
                "ShowSpeakerLabels": True,
                "MaxSpeakerLabels": DEFAULT_TRANSCRIBE_MAX_SPEAKERS,
                "ShowAlternatives": False,
                "VocabularyName": "FP-SA-Vocab-V1",
            },
        )
        print(f"Transcription job '{job_name}' started successfully.")
        return job_name
    except Exception as e:
        print(f"Error starting transcription job '{job_name}': {e}")
        raise


def calculate_openai_cost(messages, response_raw, model_name=OPENAI_MODEL) -> float:
    """
    Calculate the estimated cost of an OpenAI API call in dollars.

    Args:
        messages (list): The input messages sent to OpenAI, including roles and content.
        response_raw (dict): The raw response from OpenAI's API call.
        model_name (str): The model used for the API call. Defaults to "gpt-3.5-turbo".

    Returns:
        float: The estimated cost of the query in USD.
    """
    # Pricing per 1,000 tokens (adjust these based on OpenAI's pricing structure)
    pricing = {
        "gpt-3.5-turbo": {"input": 0.0015, "output": 0.0020},  # USD per 1,000 tokens
        "gpt-4": {"input": 0.03, "output": 0.06},             # USD per 1,000 tokens
    }

    # Tokenizer setup
    model_encodings = {
        "gpt-3.5-turbo": "cl100k_base",
        "gpt-4": "cl100k_base",
    }

    if model_name not in pricing or model_name not in model_encodings:
        raise ValueError(f"Unsupported model: {model_name}")

    # Get tokenizer
    tokenizer = tiktoken.get_encoding(model_encodings[model_name])

    # Calculate input tokens
    input_text = " ".join([msg["content"] for msg in messages])
    input_tokens = len(tokenizer.encode(input_text))

    # Calculate output tokens
    output_text = response_raw.choices[0].message.content
    output_tokens = len(tokenizer.encode(output_text))

    # Calculate total cost
    input_cost = (input_tokens / 1000) * pricing[model_name]["input"]
    output_cost = (output_tokens / 1000) * pricing[model_name]["output"]
    total_cost = input_cost + output_cost

    return total_cost


def log_query(prompt, response_msg, cost_for_query, model_name, log_file=LLM_LOG_FILE_NAME):
    with open(log_file, "a") as log:
        log_entry = {
            "timestamp": datetime.now().isoformat(),
            "model": model_name,
            "prompt": prompt,
            "response": response_msg,
            "cost_usd": round(cost_for_query, 6),
        }
        log.write(json.dumps(log_entry) + "\n")
        
        
def make_openai_query(prompt, model_name=OPENAI_MODEL, log_file=LLM_LOG_FILE_NAME):
    """
    Make a query to the OpenAI API and return the response.

    Args:
        prompt (str): The prompt for the query.
        model_name (str): The model used for the API call. Defaults to "gpt-3.5-turbo".

    Returns:
        dict: The response from the OpenAI API.
    """
    messages=[
        {"role": "system", "content": "You are a helpful assistant."},
        {"role": "user", "content": prompt}
    ]
    response = openai.chat.completions.create(
        model=model_name,
        messages=messages,
        max_tokens=int(OPENAI_MAX_TOKENS_CLASSIFICATION*1.3),
        temperature=OPENAI_TEMPERATURE,
    )
    # Extract response content
    response_msg = response.choices[0].message.content
    # Estimate Cost
    cost_for_query = calculate_openai_cost(messages=messages, response_raw=response, model_name=OPENAI_MODEL)
    cost_tracker.record_query_cost(cost_for_query)
    
    # Log the query and response
    if write_queries_to_log_file:
        log_query(prompt, response_msg, cost_for_query, model_name, log_file=log_file)
    
    return response, cost_for_query
    
    

def classify_roles_with_openai(transcript: str, n_speakers: int) -> dict:
    """Classify roles in transcript"""

    prompt = f"""You are an AI assistant tasked with classifying speakers in a transcript. 

    A speakers tag is represented by 'spk_' followed by a counter. Examples for 3 speakers: spk_0, spk_1, spk_2.
    
    Each speaker must be classified as either:
    - **"financial planner"**: The person asking questions, giving advice, or leading the discussion.
    - **"client"**: The person responding, asking for clarification, or describing their goals.

    Your output must be a JSON object which maps each speaker to their role in this exact format: {{"spk_0": "financial planner", "spk_1": "client", ...}}, where the number of items matches the number of unique speakers which is {n_speakers} speakers.

    Transcript:
    {transcript}

    Output:
    """
    response, _ = make_openai_query(prompt, model_name=OPENAI_MODEL)
    return json.loads(response.choices[0].message.content)


def summarize_transcript_with_mapping(transcript: str, speaker_mapping: dict) -> str:
    """Summarise transcript using speaker mapping."""
    prompt = f"""
    You are an AI assistant summarizing a transcript based on speaker roles. Use the following speaker mapping to identify roles: {speaker_mapping}. 

    Your summary must:
    - Be in third person and concise.
    - Clearly differentiate between what the "financial planner" (advisor) and "client" (client) said.
    - Highlight advice, questions, and goals mentioned.

    Transcript:
    {transcript}

    Summary:
    """
    response, _ = make_openai_query(prompt, model_name=OPENAI_MODEL)
    return response.choices[0].message.content

def extract_goals_with_openai(transcript: str) -> list:
    """Extracts client goals from the transcript."""
    prompt = f"""You are an AI assistant tasked with extracting goals mentioned by the client in a transcript. 

    Your output must:
    - Be a list of client goals, each as a short and specific statement.
    - Ignore any unrelated information or advice provided by the financial planner.
    - Return "None" if no hard facts are mentioned.

    Transcript:
    {transcript}

    Client Goals:
    - 
    """
    response, _ = make_openai_query(prompt, model_name=OPENAI_MODEL)
    response_msg = response.choices[0].message.content
    if response_msg.lower() == "none":
        return []
    return response_msg

def extract_hard_facts_with_openai(transcript: str) -> list:
    """Extracts hard financial facts."""
    prompt = f"""You are an AI assistant tasked with extracting hard financial facts mentioned in a transcript.

    Hard facts are:
    - Objective and verifiable financial details (e.g., account balances, tax percentages, financial commitments, or investment products mentioned).
    - Not advice, opinions, or goals.

    Your output must:
    - Be a list of financial facts.
    - Return "None" if no hard facts are mentioned.

    Transcript:
    {transcript}

    Hard Financial Facts:
    - 
    """
    response, _ = make_openai_query(prompt, model_name=OPENAI_MODEL)
    response_msg = response.choices[0].message.content
    if response_msg.lower() == "none":
        return []
    return response_msg

def extract_financial_advice_with_openai(transcript: str) -> list:
    """Extracts financial advice."""
    prompt = f"""You are an AI assistant tasked with extracting financial advice provided by the financial planner in a transcript.

    Financial advice is:
    - Specific recommendations or actionable suggestions given by the financial planner.
    - Not client goals, questions, or general statements.

    Your output must:
    - Be a list of advice statements.
    - Return "None" if no advice is mentioned.

    Transcript:
    {transcript}

    Financial Advice:
    - 
    """
    response, _ = make_openai_query(prompt, model_name=OPENAI_MODEL)
    response_msg = response.choices[0].message.content
    if response_msg.lower() == "none":
        return []
    return response_msg
  

def calculate_simple_token_count(text):
    """Estimate token count based on text length."""
    return len(text.split())  # Rough approximation: 1 word ≈ 1 token.  
    
    

def estimate_token_count(text: str, model_name: str = "gpt-3.5-turbo") -> int:
    """
    Estimate the number of tokens used by a given text for an OpenAI model.

    Args:
        text (str): The input text to estimate token usage for.
        model_name (str): The model name (e.g., "gpt-3.5-turbo", "gpt-4").
                     Different models may use different tokenizers.

    Returns:
        int: The estimated token count.
    """
    # Mapping model names to tiktoken encoders
    model_encodings = {
        "gpt-3.5-turbo": "cl100k_base",
        "gpt-4": "cl100k_base",
        "text-davinci-003": "p50k_base",
        "text-davinci-002": "p50k_base",
        "text-curie-001": "r50k_base",
        "text-babbage-001": "r50k_base",
        "text-ada-001": "r50k_base",
    }

    # Select the encoding based on the model
    if model_name not in model_encodings:
        raise ValueError(f"Unsupported model: {model_name}. Add its encoding to the mapping.")

    encoding_name = model_encodings[model_name]

    # Get the tokenizer
    try:
        tokenizer = tiktoken.get_encoding(encoding_name)
    except KeyError:
        raise ValueError(f"Encoding {encoding_name} not found. Ensure the model mapping is correct.")

    # Encode the text and return the token count
    return len(tokenizer.encode(text))


def partition_transcription_into_chunks(speaker_segment_transcriptions: list, min_tokens: int = 500, chunk_duration_s: int = 300) -> list:
    """
    Splits the transcript into chunks based on minimum token size or maximum duration.

    Parameters:
    - speaker_segment_transcriptions (list): The processed list of speaker segments.
    - min_tokens (int): Minimum number of tokens per chunk (default: 500).
    - chunk_duration_s (int): Maximum duration of a chunk in seconds (default: 300).

    Returns:
    - list: A list of chunks, where each chunk contains text and speaker information.
    """
    chunks = []
    current_chunk = []
    current_duration = 0.0
    current_token_count = 0

    def add_chunk():
        """Add the current chunk to the chunks list."""
        if current_chunk:
            chunk_text = [seg["transcript"] for seg in current_chunk]
            chunk_speakers = [{"speaker": seg["speaker_label"]} for seg in current_chunk]
            chunks.append({
                "text": chunk_text,
                "speakers": chunk_speakers,
                "start_time": current_chunk[0]["start_time"],
                "end_time": current_chunk[-1]["end_time"]
            })

    for segment in speaker_segment_transcriptions:
        # Get duration and token count for the segment
        segment_start = float(segment["start_time"])
        segment_end = float(segment["end_time"] or segment_start)
        segment_duration = segment_end - segment_start
        segment_text = segment["transcript"]
        segment_tokens = len(segment_text.split())  # Estimate token count as word count

        # Add the segment to the current chunk
        current_chunk.append(segment)
        current_duration += segment_duration
        current_token_count += segment_tokens

        # If chunk exceeds duration or token limit, finalize it
        if current_duration >= chunk_duration_s or current_token_count >= min_tokens:
            add_chunk()
            current_chunk = []
            current_duration = 0.0
            current_token_count = 0

    # Add the last chunk if any segments remain
    if current_chunk:
        add_chunk()
        
    # Now ensure the speaker label is tagged with the speaker mapping
    for k, v in speaker_mapping.items():
        for chunk in chunks:
            print(chunk)
            for seg in chunk['speakers']:
                if seg['speaker'] == k:
                    seg['speaker'] = v

    return chunks    
    
    
    
def chunk_transcript(raw_transcript: dict, min_tokens: int = 500, chunk_duration_s: int = 300) -> list:
    """
    Splits the transcript into chunks based on minimum token size or maximum duration.

    Parameters:
    - raw_transcript (dict): The raw transcription output from AWS Transcribe.
    - min_tokens (int): Minimum number of tokens per chunk (default: 500).
    - chunk_duration_s (int): Maximum duration of a chunk in seconds (default: 300).

    Returns:
    - list: A list of chunks, where each chunk contains text and speaker information.
    """
    speaker_segments = raw_transcript['speaker_labels']['segments']
    audio_segments = raw_transcript['audio_segments']

    chunks = []
    current_chunk = []
    current_duration = 0.0
    current_token_count = 0

    def add_chunk():
        """Add the current chunk to the chunks list."""
        chunk_text = "\n".join([seg['text'] for seg in current_chunk])
        chunk_speakers = [{"speaker": seg['speaker'], "text": seg['text']} for seg in current_chunk]
        chunks.append(chunk_speakers)

    for segment in speaker_segments:
        segment_start = float(segment['start_time'])
        segment_end = float(segment['end_time'])
        segment_duration = segment_end - segment_start

        # Get text for this segment
        segment_text = " ".join(
            [
                seg["transcript"]
                for seg in audio_segments
                if float(seg["start_time"]) >= float(segment["start_time"])
                and float(seg["end_time"]) <= float(segment["end_time"])
            ]
        )
        segment_tokens = estimate_token_count(segment_text)

        # Add the segment to the current chunk
        current_chunk.append({"speaker": segment['speaker_label'], "text": segment_text})
        current_duration += segment_duration
        current_token_count += segment_tokens

        # If chunk exceeds duration or token limit, finalize it
        if current_duration >= chunk_duration_s or current_token_count >= min_tokens:
            add_chunk()
            current_chunk = []
            current_duration = 0.0
            current_token_count = 0

    # Add the last chunk if any segments remain
    if current_chunk:
        add_chunk()

    return chunks
    

def parse_speaker_mapping(classification_output: dict, chunk: list) -> dict:
    """
    Parses the output of the 'Classify Roles' tool into a speaker mapping dictionary.

    Parameters:
    - classification_output (dict): The output from the "Classify Roles" tool (e.g., {'spk_0': 'financial planner', 'spk_1': 'client'}).
    - chunk (list): The list of transcript segments, where each segment contains 'speaker_label' and 'transcript'.

    Returns:
    - dict: A dictionary mapping speaker IDs (e.g., "spk_0") to their roles (e.g., "financial planner", "client").
    """
    if not isinstance(classification_output, dict):
        raise ValueError("Invalid classification output format. Expected a dictionary.")

    speaker_mapping = {}
    for segment in chunk:
        speaker = segment['speaker_label']  # e.g., "spk_0", "spk_1"
        if speaker in classification_output:
            role = classification_output[speaker]
            if speaker in speaker_mapping:
                # Check for role conflicts
                if speaker_mapping[speaker] != role:
                    print(f"Warning: Speaker '{speaker}' was previously assigned as '{speaker_mapping[speaker]}' and is now being reassigned to '{role}'.")
            else:
                speaker_mapping[speaker] = role
        else:
            print(f"Warning: Speaker '{speaker}' not found in classification output.")

    return speaker_mapping



# Initialize the tools
def get_tools() -> List[Tool]:
    return [
        Tool(
            name="Summarize Transcript",
            func=summarize_transcript_tool,
            description=(
                "Summarize the transcript by identifying key insights and main points. "
                "Input should be a JSON string with 'transcript' and 'speaker_mapping'."
            ),
        ),
        Tool(
            name="Extract Hard Facts",
            func=extract_hard_facts_tool,
            description=(
                "Extract hard, factual information mentioned in the transcript. "
                "Input should be the transcript as a single string."
            ),
        ),
        Tool(
            name="Extract Financial Advice",
            func=extract_financial_advice_tool,
            description=(
                "Extract financial advice or actionable recommendations from the transcript. "
                "Input should be the transcript as a single string."
            ),
        ),
    ]

# Define your tool functions with a single input argument

def summarize_transcript_tool(input_json: str) -> str:
    """
    Summarizes the transcript with the given speaker mapping.
    Expects a JSON string with 'transcript' and 'speaker_mapping'.
    """
    try:
        input_data = json.loads(input_json)
        transcript = input_data['transcript']
        speaker_mapping = input_data['speaker_mapping']
        summary = summarize_transcript_with_mapping(transcript=transcript, speaker_mapping=speaker_mapping)
        return summary
    except (json.JSONDecodeError, KeyError) as e:
        print(f"Invalid input for summarize_transcript_tool: {e}")
        return "Error: Invalid input format for summarizing transcript."

def extract_hard_facts_tool(transcript: str) -> str:
    """
    Extracts hard facts from the transcript.
    """
    try:
        facts = extract_hard_facts_with_openai(transcript)
        return facts
    except Exception as e:
        print(f"Error in extract_hard_facts_tool: {e}")
        return "Error: Failed to extract hard facts."


def extract_financial_advice_tool(transcript: str) -> str:
    """
    Extracts financial advice from the transcript.
    """
    try:
        advice = extract_financial_advice_with_openai(transcript)
        return advice
    except Exception as e:
        print(f"Error in extract_financial_advice_tool: {e}")
        return "Error: Failed to extract financial advice."


# Initialize the agent once
def create_agent() -> Optional:
    tools = get_tools()

    try:
        llm = OpenAI(temperature=0, max_tokens=1000)
        agent = initialize_agent(
            tools=tools,
            llm=llm,
            agent="zero-shot-react-description",
            verbose=True
        )
        print("Agent initialized successfully.")
        return agent
    except Exception as e:
        print(f"Failed to initialize the agent: {e}")
        return None
    
def process_audio_segment(chunk: List[Dict], speaker_mapping: Dict[str, str], agent) -> Dict[str, Optional[str]]:
    """
    Processes a single transcript chunk using the initialized agent to summarize, extract hard facts, and extract advice.

    Parameters:
    - chunk (list): A chunk of transcript segments.
    - speaker_mapping (dict): A dictionary mapping speakers to roles.
    - agent: The initialized Langchain agent.

    Returns:
    - dict: Processed results for the chunk, including summary, hard facts, and advice.
    """
    # Combine chunk into a single transcript text    
    pre_all_chunk_text = []
    for seg_text, speaker in zip(chunk['text'], chunk['speakers']):
        pre_all_chunk_text.append(f"{speaker['speaker']}: {seg_text}")
    
    all_chunk_text = "\n".join(pre_all_chunk_text)
    print(all_chunk_text)
        
    # Prepare JSON inputs for each tool
    summarize_input = json.dumps({
        "transcript": all_chunk_text,
        "speaker_mapping": speaker_mapping
    })
    
    hard_facts_input = json.dumps({
        "transcript": all_chunk_text
    })
    
    financial_advice_input = json.dumps({
        "transcript": all_chunk_text
    })
    # Define the tool inputs
    tool_inputs = {
        "Summarize Transcript": summarize_input,
        "Extract Hard Facts": hard_facts_input,
        "Extract Financial Advice": financial_advice_input,
    }

    
    results = {}
    # Iterate over each tool and execute its function with appropriate arguments
    for tool in agent.tools:
        try:
            input_for_tool = tool_inputs.get(tool.name, "")
            # Call other tools with just the transcript
            tool_result = tool.run(input_for_tool)
            
            # Store the result with a formatted key
            results[tool.name.lower().replace(" ", "_")] = tool_result
            print(f"Tool '{tool.name}' executed successfully.")
        except Exception as e:
            print(f"Error running tool '{tool.name}': {e}")
            results[tool.name.lower().replace(" ", "_")] = None

    return results



def generate_combined_summary(summaries: list) -> str:
    """
    Combines individual summaries into a cohesive final summary using OpenAI.

    Parameters:
    - summaries (list): List of individual chunk summaries.

    Returns:
    - str: A single cohesive summary.
    """
    summaries_text = "\n\n".join(summaries)
    # prompt = f"""Combine the following summaries into a cohesive, comprehensive summary of a few sentences long. Text to summarize:

    # {summaries_text}
    
    # Ensure the final summary is concise, clear, and captures all key points from the individual summaries.
    # """
    prompt = f"""You are an AI assistant tasked with combining multiple summaries into one cohesive and concise summary. 

    Your goal is to:
    - Combine the following summaries into a single, comprehensive summary that captures all the key points.
    - Ensure the final summary is clear, concise, and only a few sentences long.

    Individual summaries:
    {summaries_text}

    Final cohesive summary:
    """
    response, cost_for_query = make_openai_query(prompt, model_name=OPENAI_MODEL)
    response_msg = response.choices[0].message.content
    return response_msg, cost_for_query

def generate_combined_advice(advice_list: list) -> str:
    """
    Combines advice into a concise, deduplicated bullet-point list using OpenAI.

    Parameters:
    - advice_list (list): List of advice extracted from chunks.

    Returns:
    - str: Deduplicated and consolidated advice.
    """
    advice_text = "\n".join(advice_list)

    prompt = f"""You are an AI assistant tasked with deduplicating and consolidating financial advice into a clear, organized list.

    Your goal is to:
    - Combine the following advice points into a concise and actionable bullet-point list.
    - Deduplicate similar advice and ensure there is no overlap or redundancy.
    - Return the advice in the format: ['Advice 1', 'Advice 2', 'Advice 3'].
    - If there is no clear advice, return 'None'.

    Advice points to consolidate:
    {advice_text}

    Final consolidated advice:
    """
    response, cost_for_query = make_openai_query(prompt, model_name=OPENAI_MODEL)
    response_msg = response.choices[0].message.content
    if response_msg.lower() == "none":
        return []
    return response_msg, cost_for_query

def generate_combined_hard_facts(hard_facts_list: list) -> str:
    """
    Combines hard facts into a concise, deduplicated bullet-point list using OpenAI.

    Parameters:
    - hard_facts_list (list): List of hard facts extracted from chunks.

    Returns:
    - str: Deduplicated and consolidated hard facts as a list.
    """
    hard_facts_text = "\n".join(hard_facts_list)

    prompt = f"""You are an AI assistant tasked with consolidating and deduplicating a list of hard facts. 
    A hard fact is a statement of objective truth, free from ambiguity, interpretation, or subjectivity.

    Your goal is to:
    - Extract only hard facts from the provided list, avoiding any overlap or redundancy.
    - Ensure each fact is unique and does not repeat information stated elsewhere in the list.
    - Exclude any subjective opinions, recommendations, or ambiguous statements.
    - Present the final output in the format: ['Fact 1', 'Fact 2', 'Fact 3'].
    - If no valid hard facts can be extracted, return 'None'.

    Hard facts to analyze and consolidate:
    {hard_facts_text}

    Final consolidated hard facts:
    """
    
    response, cost_for_query = make_openai_query(prompt, model_name=OPENAI_MODEL)
    response_msg = response.choices[0].message.content
    if response_msg.lower() == "none":
        return []
    return response_msg, cost_for_query


def aggregate_chunk_results(chunk_results: list) -> dict:
    """
    Aggregates results from all processed chunks into a final summary, list of hard facts, 
    and optionally goals and advice.

    Parameters:
    - chunk_results (list): List of dictionaries containing processed results for each chunk.

    Returns:
    - dict: Aggregated results including final summary, hard facts, goals, and advice.
    """
    # Extract and prepare data for aggregation
    summaries = [chunk["summarize_transcript"] for chunk in chunk_results]
    all_advice = [chunk["extract_financial_advice"] for chunk in chunk_results if chunk["extract_financial_advice"]!= []]
    all_hard_facts = [chunk["extract_hard_facts"] for chunk in chunk_results]
    # all_goal = [chunk["extract_goals"] for chunk in chunk_results]
    
    # Generate consolidated outputs
    final_full_summary, _ = generate_combined_summary(summaries)
    final_combined_advice, _ = generate_combined_advice(all_advice)  # Optional, can be uncommented later
    final_combined_hard_facts, _ = generate_combined_hard_facts(all_hard_facts)  # Optional, can be uncommented later

    return {
        "final_full_summary": final_full_summary,
        "final_combined_advice": final_combined_advice,
        "all_hard_facts": final_combined_hard_facts,
        # "all_goals": all_goals,
    }



def generate_speaker_mapping(combined_speaker_segments: List[dict]) -> dict:
    """
    Generates a speaker mapping by classifying roles in the entire transcript.

    Parameters:
    - raw_transcript (dict): The raw transcription output from AWS Transcribe.

    Returns:
    - dict: A dictionary mapping speaker IDs to their roles (e.g., "financial planner", "client").
    """
    # Combine all segments into a single transcript text
    transcript_text = "\n".join(
        [f"{seg['speaker_label']}: {seg['transcript']}" for seg in combined_speaker_segments]
    )
    n_speakers = len(set(seg['speaker_label'] for seg in combined_speaker_segments))
    print(f"n_speakers: {n_speakers}")

    # Classify roles using OpenAI
    classification_output = classify_roles_with_openai(transcript=transcript_text, n_speakers=n_speakers)

    # Parse the classification output to generate the speaker mapping
    speaker_mapping = parse_speaker_mapping(classification_output, combined_speaker_segments)
    return speaker_mapping
    
    

def create_transcript_output_docx(combined_speaker_segments: List[dict], speaker_mapping: dict, meeting_name: str, fp_name="Morgan", client_name: str = "Client A"):
    """
    Create a formatted DOCX file from a AWS Transcribe output.

    Args:
        combined_speaker_segments (dict): List of Speaker Segments with speaker and transcript as dict keys
        speaker_mapping (dict): Mapped dictionary of speaker indices to speaker roles
        client_name (str): Name of the client.
        meeting_name (str): Name of the meeting.
        fp_name (str): Name of the FP (default is "Morgan").
    """
    # Initialize document
    document = Document()

    # Add metadata section
    document.add_heading(f"Client: {client_name}", level=1)
    document.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    document.add_paragraph(f"Meeting Name: {meeting_name}")
    document.add_paragraph(f"Financial Planner: {fp_name}")
    document.add_paragraph("_" * 50)
    document.add_paragraph("Speaker Mapping:")
    for speaker_label, speaker_name in speaker_mapping.items():
        document.add_paragraph(f"{chr(ord('A') + list(speaker_mapping.keys()).index(speaker_label))}: {speaker_name}")

    document.add_paragraph("\n")
    document.add_paragraph("Transcript")
    document.add_paragraph("_" * 50)

    # Add combined transcript to the document
    for segment in combined_speaker_segments:
        speaker_label = segment.get('speaker_label')
        transcript = segment.get('transcript')
        speaker_name = chr(ord('A') + list(speaker_mapping.keys()).index(speaker_label))  # Map spk_0, spk_1, etc. to A, B, C...
        document.add_paragraph(f"{speaker_name}: {transcript}")

    # Save the document
    filename = os.path.join(
        transcript_out_dir,
        f"transcript_{meeting_name}_{datetime.now().strftime('%d-%m-%y')}.docx"
    )
    document.save(filename)
    print(f"Transcript saved as {filename}")
    
    
def create_combined_speaker_segments(raw_transcript: dict) -> List[dict]:
    # Combine consecutive segments by the same speaker
    combined_segments = []
    previous_speaker = None
    current_transcript = ""

    for segment in raw_transcript.get('audio_segments', {}):
        speaker_label = segment.get('speaker_label')
        transcript = segment.get('transcript')

        if speaker_label == previous_speaker:
            current_transcript += f" {transcript}"  # Append to current transcript
        else:
            if previous_speaker is not None:
                combined_segments.append({"speaker_label": previous_speaker, "transcript": current_transcript})
            previous_speaker = speaker_label
            current_transcript = transcript

    # Add the last segment
    if previous_speaker is not None:
        combined_segments.append({"speaker_label": previous_speaker, "transcript": current_transcript})
        
    return combined_segments



def merge_transcript_speaker_breaks(raw_transcripts: dict) -> dict:
    """
    Merges consecutive segments in the AWS Transcribe output where the same speaker is speaking,
    updating timestamps and preserving the format.
    
    :param raw_transcripts: Dictionary containing AWS Transcribe raw output.
    :return: Updated dictionary with merged speaker segments.
    """
    merged_audio_segments = []
    current_segment = None
    current_items = []
    items = raw_transcripts["items"]

    # Ensure all items have an end_time by inferring from the next item's start_time
    for i in range(len(items) - 1):
        if "end_time" not in items[i]:
            items[i]["end_time"] = items[i + 1]["start_time"]
    if "end_time" not in items[-1]:  # Handle the last item
        items[-1]["end_time"] = items[-1].get("start_time")  # Keep it as start_time if nothing else

    for segment in raw_transcripts["audio_segments"]:
        # Start a new segment if it's the first or speaker changes
        if not current_segment or segment.get("speaker_label") != current_segment.get("speaker_label"):
            # Save the completed segment
            if current_segment:
                current_segment["start_time"] = current_items[0]["start_time"]
                current_segment["end_time"] = current_items[-1]["end_time"]
                current_segment["transcript"] = " ".join(
                    [item["alternatives"][0]["content"] for item in current_items]
                )
                current_segment["items"] = [item["id"] for item in current_items]
                merged_audio_segments.append(current_segment)

            # Start a new segment
            current_segment = {
                "id": len(merged_audio_segments),
                "speaker_label": segment.get("speaker_label"),
                "start_time": None,
                "end_time": None,
                "transcript": "",
                "items": []
            }
            current_items = []

        # Collect items for the current segment
        for item_id in segment["items"]:
            current_items.append(items[item_id])

    # Add the final segment
    if current_segment:
        current_segment["start_time"] = current_items[0]["start_time"]
        current_segment["end_time"] = current_items[-1]["end_time"]
        current_segment["transcript"] = " ".join(
            [item["alternatives"][0]["content"] for item in current_items]
        )
        current_segment["items"] = [item["id"] for item in current_items]
        merged_audio_segments.append(current_segment)

    # Update the raw_transcripts dictionary
    raw_transcripts["audio_segments"] = merged_audio_segments

    # Concatenate transcripts into logical sections for the 'transcripts' key
    concatenated_transcripts = []
    current_transcript = ""

    for segment in merged_audio_segments:
        # Append the current segment's transcript
        if current_transcript:
            current_transcript += " " + segment["transcript"]
        else:
            current_transcript = segment["transcript"]

        # Logic to end a concatenation group (based on speaker changes or other conditions)
        if "speaker_label" in segment:
            concatenated_transcripts.append(current_transcript)
            current_transcript = ""

    # Add the final group if any
    if current_transcript:
        concatenated_transcripts.append(current_transcript)

    # Update the transcripts key
    raw_transcripts["transcripts"] = [{"transcript": t} for t in concatenated_transcripts]
    return raw_transcripts



def extract_confidence_statistics(raw_transcription: dict):
    """
    Extracts confidence scores, calculates averages and standard deviations for:
    - 5-minute intervals
    - Total transcription confidence
    - Confidence by speaker

    Parameters:
    - raw_transcription (dict): Raw AWS Transcribe data.

    Returns:
    - dict: Statistics with confidence scores in 5-minute intervals, overall, and per speaker.
    """
    # Extract items and speaker information
    items = raw_transcription["items"]
    speaker_segments = raw_transcription["speaker_labels"]["segments"]

    # Collect confidence scores with timestamps
    confidence_scores = []
    for item in items:
        if "alternatives" in item and "confidence" in item["alternatives"][0]:
            start_time = float(item["start_time"]) if "start_time" in item else None
            end_time = float(item["end_time"]) if "end_time" in item else None
            confidence = float(item["alternatives"][0]["confidence"])
            confidence_scores.append({"start_time": start_time, "end_time": end_time, "confidence": confidence})

    # Calculate overall confidence statistics
    confidences = [score["confidence"] for score in confidence_scores]
    total_avg_confidence = np.mean(confidences)
    total_std_confidence = np.std(confidences)

    # 5-minute interval confidence statistics
    interval_confidences = []
    interval_stats = {}
    current_interval = 0
    for score in confidence_scores:
        interval_start = current_interval * 300
        interval_end = (current_interval + 1) * 300

        if score["start_time"] and interval_start <= score["start_time"] < interval_end:
            interval_confidences.append(score["confidence"])
        elif score["start_time"] and score["start_time"] >= interval_end:
            # Process the completed interval
            if interval_confidences:
                avg_confidence = np.mean(interval_confidences)
                std_confidence = np.std(interval_confidences)
                interval_stats[f"Interval {current_interval}"] = {"avg": avg_confidence, "std": std_confidence}
                interval_confidences = []
            # Move to the next interval
            current_interval += 1
            if current_interval * 300 <= score["start_time"] < (current_interval + 1) * 300:
                interval_confidences.append(score["confidence"])

    # Add the last interval if not empty
    if interval_confidences:
        avg_confidence = np.mean(interval_confidences)
        std_confidence = np.std(interval_confidences)
        interval_stats[f"Interval {current_interval}"] = {"avg": avg_confidence, "std": std_confidence}

    # Speaker-specific confidence statistics
    speaker_confidences = {}
    for segment in speaker_segments:
        speaker_label = segment["speaker_label"]
        segment_start = float(segment["start_time"])
        segment_end = float(segment["end_time"])
        segment_scores = [
            score["confidence"]
            for score in confidence_scores
            if score["start_time"] and segment_start <= score["start_time"] < segment_end
        ]
        if speaker_label not in speaker_confidences:
            speaker_confidences[speaker_label] = []
        speaker_confidences[speaker_label].extend(segment_scores)

    speaker_stats = {
        speaker: {
            "avg": np.mean(scores),
            "std": np.std(scores),
            "count": len(scores),
        }
        for speaker, scores in speaker_confidences.items()
        if scores
    }

    # Print results
    print("Total Transcription Confidence:")
    print(f"Average: {total_avg_confidence:.2f}, Standard Deviation: {total_std_confidence:.2f}\n")
    print("5-Minute Interval Confidence:")
    for interval, stats in interval_stats.items():
        print(f"{interval}: Avg={stats['avg']:.2f}, Std={stats['std']:.2f}")
    print("\nSpeaker Confidence:")
    for speaker, stats in speaker_stats.items():
        print(f"{speaker}: Avg={stats['avg']:.2f}, Std={stats['std']:.2f}, Count={stats['count']}")

    return {
        "total_confidence": {"avg": total_avg_confidence, "std": total_std_confidence},
        "interval_confidence": interval_stats,
        "speaker_confidence": speaker_stats,
    }

if __name__ == "__main__":

    # Step 1: Initialize clients and managers
    boto3_client, s3_manager, transcription_manager = initialize_clients()
    print(f"Initialized Client Connections Successfully")
    if not DRY_RUN_SYSTEM:
        # Step 2: Upload audio
        s3_file_uri = upload_audio_to_s3(s3_manager, "uploads/FP Chat Morgan - 19 Jan 2025/Partial-1-18min-Audio-Clip-Ryan-Mccarlie.wav")
        print(f"Uploaded S3 file uri: {s3_file_uri}")
        
        try:
            # Step 3: Start transcription job
            job_name = start_aws_transcription_job(
                transcription_manager=transcription_manager,
                input_s3_uri=s3_file_uri,
                output_s3_dir_prefix="other/audio-files/output/",
            )
            print(f"Job name: {job_name}")
            
            # Step 4: Wait for transcription to complete
            response = transcription_manager.wait_for_transcription_completion(job_name)
        
        except Exception as e:
            print(f"Error starting transcription job: {e}")
            raise e

    else:
        job_name = "transcription_job_20250120_221929"

    # Step 5: Fetch transcription output
    raw_transcripts = transcription_manager.fetch_transcription_output(job_name)
    print(f"Successfully fetched raw transcriptions")
    
    confidence_stats = extract_confidence_statistics(raw_transcription=raw_transcripts)
    print("confidence_stats", confidence_stats)
    
    merged_transcripts = merge_transcript_speaker_breaks(raw_transcripts=raw_transcripts)
    combined_speaker_segments = merged_transcripts["audio_segments"]
    print(combined_speaker_segments)
    # combined_speaker_segments = create_combined_speaker_segments(raw_transcript=raw_transcripts)
    # print(combined_speaker_segments)
    
    # Generate Speaker Mapping like so
    # speaker_mapping = generate_speaker_mapping(combined_speaker_segments=combined_speaker_segments)
    speaker_mapping = {'spk_0': 'financial planner', 'spk_1': 'client'}
    print("speaker_mapping:\n", speaker_mapping)    
    
    meeting_name = "Intro FP Chat with Ryan"
    client_name = "Ryan"
    create_transcript_output_docx(
        combined_speaker_segments=combined_speaker_segments, 
        speaker_mapping=speaker_mapping,
        meeting_name=meeting_name, client_name=client_name, fp_name=DEFAULT_FP_NAME
    )

    # Step 2: Chunk the transcript into manageable pieces
    transcription_chunks = partition_transcription_into_chunks(speaker_segment_transcriptions=combined_speaker_segments, chunk_duration_s=300)
    print(f"Number of transcription_chunks produced: {len(transcription_chunks)}")
    print(f"Size of first chunk: {len(transcription_chunks[0])}")
    print(f"First chunk: {transcription_chunks[0]}")    
    print(f"Second chunk: {transcription_chunks[1]}")
        
    # Initialize tools and agent
    agent = create_agent()

    if agent is None:
        raise RuntimeError("Agent initialization failed. Cannot process audio segments.")
    
    # Step 3: Process each chunk    
    processed_audio_results = []
    for audio_segments in transcription_chunks:
        processed_results = process_audio_segment(chunk=audio_segments, speaker_mapping=speaker_mapping, agent=agent)
        processed_audio_results.append(processed_results)

    # Step 4: Aggregate the results from all chunks
    final_results = aggregate_chunk_results(processed_audio_results)

    # # Output final result
    print(json.dumps(final_results, indent=4))
    
    # Get full query stats
    stats = cost_tracker.get_query_stats()
    print("Query Statistics:")
    for key, value in stats.items():
        print(f"{key}: {value}")
    
    # Log the final results
    log_final_results(final_results, stats)
