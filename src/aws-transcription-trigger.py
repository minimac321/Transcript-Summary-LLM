import openai
import os
import json
import re
import time
import sys
import numpy as np
from datetime import datetime
from dotenv import load_dotenv
from openai import OpenAI as OpenAIClient
import google.generativeai as genai
from vertexai.preview.generative_models import GenerativeModel
from vertexai import init

from docx import Document
from typing import Dict, List, Optional, Tuple

from langchain.agents import initialize_agent, Tool
from langchain_community.llms import OpenAI as LangChainOpenAI 
from langchain.agents import AgentExecutor

from aws_utils import initialize_clients, start_aws_transcription_job, upload_audio_to_s3
from constants import DEFAULT_FP_NAME, GCP_LOCATION, GCP_PROJECT_ID, GEMINI_MODEL_NAMES_LIST, GEMINI_PRICING_DICT, MODEL_TOKEN_LIMITS, OPENAI_MODEL_NAMES_LIST, OPENAI_ORG_NAME, OPENAI_PRICING_DICT, OPENAI_TRANSCRIPTION_PROJ_ID, USD_TO_ZAR_CONVERSION
from gemini_llm_functions import SpeakerClassification, SpeakerLabel, make_gemini_query
from log_utils import log_final_results
from openai_llm_functions import make_openai_query
from utils import count_tokens_func, split_text_into_chunks
from cost_tracker import QueryCostTracker

# Load .env variables
load_dotenv("app/.env")

# Initialize OpenAI API key (replace with your key)
openai.api_key = os.getenv("OPENAI_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# OPENAI_TRANSCRIPTION_PROJ_NAME = "Transcription-LLM"

openai_client = OpenAIClient(
    api_key=OPENAI_API_KEY,
    organization=OPENAI_ORG_NAME,
    project=OPENAI_TRANSCRIPTION_PROJ_ID,
)

# AWS Credentials from environment variables
AWS_ACCESS_KEY_ID = os.getenv("AWS_ACCESS_KEY_ID")
AWS_SECRET_ACCESS_KEY = os.getenv("AWS_SECRET_ACCESS_KEY")

# Initialize Gemini API
genai.configure(api_key=os.environ["GEMINI_API_KEY"])
init(project=GCP_PROJECT_ID, location=GCP_LOCATION)  # Set the correct region


LOG_DIR = "logs"
LLM_LOG_FILE_NAME = f"{LOG_DIR}/llm_api_query_logs.txt"
# LLM_LOG_FILE_NAME = f"{LOG_DIR}/gemini_api_log.txt"

today_date = datetime.now().strftime("%d_%b")
FULL_APP_LOG_FILE_NAME = f"{LOG_DIR}/full_app_output_{today_date}.txt"

outputs_dir = "outputs"
transcript_out_dir = os.path.join(outputs_dir, "transcripts")


cost_tracker = QueryCostTracker(
    usd_to_zar_conversion_rate=USD_TO_ZAR_CONVERSION,
    nth_decimal_point=8
)

#########################
### SCRIPT PARAMETERS ###
#########################
DRY_RUN_SYSTEM = True
SELECTED_MODEL = "gemini-1.5-flash"


def make_llm_query(prompt: str, model_name: str, log_file: str) -> Tuple[str, float]:
    if model_name in GEMINI_MODEL_NAMES_LIST:
        response, cost_for_query = make_gemini_query(
            prompt=prompt, model_name=model_name, 
            cost_tracker=cost_tracker,
            gemini_pricing_dict=GEMINI_PRICING_DICT, log_file=log_file,
        )
        return response, cost_for_query
    
    elif model_name in OPENAI_MODEL_NAMES_LIST:
        response, cost_for_query = make_openai_query(
            prompt=prompt, model_name=model_name, 
            cost_tracker=cost_tracker,
            openai_pricing_dict=OPENAI_PRICING_DICT, log_file=log_file,
        )
        return response, cost_for_query
    
    else:
        print("(ERROR) - Model is not part of OpenAI or Gemini models")
        assert False

def classify_speakers_api_call(n_speakers: int, transcript: str):
    prompt = f"""You are an AI assistant tasked with classifying speakers in a transcript.

    A speaker's tag is represented by 'spk_' followed by a counter. Examples for 3 speakers: spk_0, spk_1, spk_2.

    Each speaker must be classified ONCE as either:
    - **{SpeakerLabel.FINANCIAL_PLANNER.value}**: The person asking questions, giving advice, or leading the discussion.
    - **{SpeakerLabel.CLIENT.value}**: The person responding, asking for clarification, or describing their goals.

    Your response must be **valid JSON only** with no extra text or formatting.  
    The output format **must match this exactly**:
    {{
        "spk_0": {SpeakerLabel.FINANCIAL_PLANNER.value},
        "spk_1": {SpeakerLabel.CLIENT.value},
        "spk_2": {SpeakerLabel.CLIENT.value}
    }}
    Ensure that the number of entries matches exactly {n_speakers} unique speakers.

    ### **Transcript:**
    {transcript}

    ### **Output (Strict JSON Format Only, No Extra Text):**
    """
    model = genai.GenerativeModel(SELECTED_MODEL)
    response = model.generate_content(
        contents=prompt,
        generation_config={
            'response_mime_type': 'application/json',
        },
    )
    # Parse JSON and assert correctness
    try:
        speaker_mapping = json.loads(response.text.strip())  # Clean and parse JSON
        
        # Validate keys and values
        pattern = r'^spk_\d+$'
        assert all(
            re.match(pattern, key) and value in {SpeakerLabel.FINANCIAL_PLANNER.value, SpeakerLabel.CLIENT.value} 
            for key, value in speaker_mapping.items()
        )
        return speaker_mapping
    except json.JSONDecodeError as e:
        print("Error parsing JSON response:", e)
        return None
    except Exception as e:
        print("Validation error:", e)
        return None


def summarize_transcript_with_mapping(transcript: str, speaker_mapping: dict) -> str:
    """Summarizes a transcript using speaker mapping to distinguish roles and key insights."""

    prompt = f"""
    You are an AI assistant summarizing a financial discussion based on speaker roles. Use the following speaker mapping to correctly identify each role: {speaker_mapping}. 

    ### **Instructions:**
    - **Write in third person** with a professional, neutral tone.
    - **Clearly differentiate** between what the **financial planner (advisor)** and the **client** said.
    - **Summarize the key points** of the conversation, focusing on:
      - **Client's financial goals** (e.g., savings, investments, debt plans).
      - **Advice provided by the financial planner** (recommendations, strategies, clarifications).
      - **Important questions or concerns raised** by the client.
    - **Avoid unnecessary details** and keep the summary concise and well-structured.
    
    ### **Transcript:**
    {transcript}

    ### **Summarized Conversation:**
    """ 
    response, _ = make_llm_query(prompt=prompt, model_name=SELECTED_MODEL, log_file=LLM_LOG_FILE_NAME)
    return response

def extract_goals_with_llm(transcript: str) -> list:
    """Extracts client-stated financial goals from the transcript while ignoring planner advice or general statements."""

    prompt = f"""You are an AI assistant tasked with extracting **clear, specific financial goals** directly mentioned by the client in a transcript.

    A **financial goal** is:
    - A **specific, actionable objective** the client intends to achieve (e.g., "Save $20,000 for a home" or "Pay off my student loans").
    - **Personal to the client**, reflecting their aspirations, priorities, or long-term financial plans.
    - **Not general financial advice, planner recommendations, or industry trends.**
    - **Not vague statements**—goals must be well-defined.

    Your task:
    - **Extract only the client's goals** (not the financial planner's advice or general discussions).
    - **Ensure goals are concise and properly formatted as a list of goals. For Example: ["Goal 1", "Goal 2", ..].**
    - **Ignore unrelated information, financial advice, or general facts.**
    - **Return "None" if no client goals are mentioned.**

    ### **Transcript:**
    {transcript}

    ### **Extracted Client Goals:**
    - 
    """
    response_msg, _ = make_llm_query(prompt=prompt, model_name=SELECTED_MODEL, log_file=LLM_LOG_FILE_NAME)
    if response_msg.lower() == "none":
        return []
    return response_msg

def extract_hard_facts_with_llm(transcript: str) -> list:
    """Extracts hard financial facts while ensuring general information not specific to the client's situation is excluded."""
    
    prompt = f"""You are an AI assistant tasked with extracting **hard financial facts** mentioned in a transcript.

    Hard facts are:
    - **Objective and verifiable financial details** directly linked to the client's financial or personal situation.
    - Includes **account balances, specific tax details affecting the client, investment products owned or considered, financial commitments, or precise monetary figures.**
    - **Does not include general financial knowledge, tax laws, or statistics that are not specific to the client.**
    - Not advice, opinions, or financial goals.

    Your output must:
    - Be a **list of hard financial facts**.
    - Exclude any general financial knowledge that does not relate to the client.
    - Return "None" if no relevant hard facts are mentioned.

    ### **Transcript:**
    {transcript}

    ### **Extracted Hard Financial Facts:**
    - 
    """
    response_msg, _ = make_llm_query(prompt=prompt, model_name=SELECTED_MODEL, log_file=LLM_LOG_FILE_NAME)
    if response_msg.lower() == "none":
        return []
    return response_msg

def extract_financial_advice_with_llm(transcript: str) -> list:
    """Extracts financial advice."""
    prompt = f"""You are an AI assistant tasked with extracting financial advice provided by the financial planner in a transcript.

    Financial advice is:
    - Specific recommendations or actionable suggestions given by the financial planner.
    - Not client goals, questions, or general statements.

    Your output must:
    - Be a list of advice statements.
    - Return "None" if no advice is mentioned.

    Transcript:
    {transcript}

    Financial Advice:
    - 
    """
    response_msg, _ = make_llm_query(prompt=prompt, model_name=SELECTED_MODEL, log_file=LLM_LOG_FILE_NAME)
    if response_msg.lower() == "none":
        return []
    return response_msg
  

def calculate_simple_token_count(text):
    """Estimate token count based on text length."""
    return len(text.split())  # Rough approximation: 1 word ≈ 1 token.  
    
    

def estimate_token_count(text: str, model_name: str = "gpt-3.5-turbo") -> int:
    """
    Estimate the number of tokens used by a given text for an OpenAI model.

    Args:
        text (str): The input text to estimate token usage for.
        model_name (str): The model name (e.g., "gpt-3.5-turbo", "gpt-4").
                     Different models may use different tokenizers.

    Returns:
        int: The estimated token count.
    """
    # Mapping model names to tiktoken encoders
    model_encodings = {
        "gpt-3.5-turbo": "cl100k_base",
        "gpt-4": "cl100k_base",
        "text-davinci-003": "p50k_base",
        "text-davinci-002": "p50k_base",
        "text-curie-001": "r50k_base",
        "text-babbage-001": "r50k_base",
        "text-ada-001": "r50k_base",
    }

    # Select the encoding based on the model
    if model_name not in model_encodings:
        raise ValueError(f"Unsupported model: {model_name}. Add its encoding to the mapping.")

    encoding_name = model_encodings[model_name]

    # Get the tokenizer
    try:
        tokenizer = tiktoken.get_encoding(encoding_name)
    except KeyError:
        raise ValueError(f"Encoding {encoding_name} not found. Ensure the model mapping is correct.")

    # Encode the text and return the token count
    return len(tokenizer.encode(text))


def partition_transcription_into_chunks(speaker_segment_transcriptions: list, min_tokens: int = 500, chunk_duration_s: int = 300) -> list:
    """
    Splits the transcript into chunks based on minimum token size or maximum duration.

    Parameters:
    - speaker_segment_transcriptions (list): The processed list of speaker segments.
    - min_tokens (int): Minimum number of tokens per chunk (default: 500).
    - chunk_duration_s (int): Maximum duration of a chunk in seconds (default: 300).

    Returns:
    - list: A list of chunks, where each chunk contains text and speaker information.
    """
    chunks = []
    current_chunk = []
    current_duration = 0.0
    current_token_count = 0

    def add_chunk():
        """Add the current chunk to the chunks list."""
        if current_chunk:
            chunk_text = [seg["transcript"] for seg in current_chunk]
            chunk_speakers = [{"speaker": seg["speaker_label"]} for seg in current_chunk]
            chunks.append({
                "text": chunk_text,
                "speakers": chunk_speakers,
                "start_time": current_chunk[0]["start_time"],
                "end_time": current_chunk[-1]["end_time"]
            })

    for segment in speaker_segment_transcriptions:
        # Get duration and token count for the segment
        segment_start = float(segment["start_time"])
        segment_end = float(segment["end_time"] or segment_start)
        segment_duration = segment_end - segment_start
        segment_text = segment["transcript"]
        segment_tokens = len(segment_text.split())  # Estimate token count as word count

        # Add the segment to the current chunk
        current_chunk.append(segment)
        current_duration += segment_duration
        current_token_count += segment_tokens

        # If chunk exceeds duration or token limit, finalize it
        if current_duration >= chunk_duration_s or current_token_count >= min_tokens:
            add_chunk()
            current_chunk = []
            current_duration = 0.0
            current_token_count = 0

    # Add the last chunk if any segments remain
    if current_chunk:
        add_chunk()
        
    # Now ensure the speaker label is tagged with the speaker mapping
    for k, v in speaker_mapping.items():
        for chunk in chunks:
            for seg in chunk['speakers']:
                if seg['speaker'] == k:
                    seg['speaker'] = v

    return chunks    
    
    
    
def chunk_transcript(raw_transcript: dict, min_tokens: int = 500, chunk_duration_s: int = 300) -> list:
    """
    Splits the transcript into chunks based on minimum token size or maximum duration.

    Parameters:
    - raw_transcript (dict): The raw transcription output from AWS Transcribe.
    - min_tokens (int): Minimum number of tokens per chunk (default: 500).
    - chunk_duration_s (int): Maximum duration of a chunk in seconds (default: 300).

    Returns:
    - list: A list of chunks, where each chunk contains text and speaker information.
    """
    speaker_segments = raw_transcript['speaker_labels']['segments']
    audio_segments = raw_transcript['audio_segments']

    chunks = []
    current_chunk = []
    current_duration = 0.0
    current_token_count = 0

    def add_chunk():
        """Add the current chunk to the chunks list."""
        chunk_text = "\n".join([seg['text'] for seg in current_chunk])
        chunk_speakers = [{"speaker": seg['speaker'], "text": seg['text']} for seg in current_chunk]
        chunks.append(chunk_speakers)

    for segment in speaker_segments:
        segment_start = float(segment['start_time'])
        segment_end = float(segment['end_time'])
        segment_duration = segment_end - segment_start

        # Get text for this segment
        segment_text = " ".join(
            [
                seg["transcript"]
                for seg in audio_segments
                if float(seg["start_time"]) >= float(segment["start_time"])
                and float(seg["end_time"]) <= float(segment["end_time"])
            ]
        )
        segment_tokens = estimate_token_count(segment_text)

        # Add the segment to the current chunk
        current_chunk.append({"speaker": segment['speaker_label'], "text": segment_text})
        current_duration += segment_duration
        current_token_count += segment_tokens

        # If chunk exceeds duration or token limit, finalize it
        if current_duration >= chunk_duration_s or current_token_count >= min_tokens:
            add_chunk()
            current_chunk = []
            current_duration = 0.0
            current_token_count = 0

    # Add the last chunk if any segments remain
    if current_chunk:
        add_chunk()

    return chunks


# Initialize the tools
def get_tools() -> List[Tool]:
    return [
        Tool(
            name="Summarize Transcript",
            func=summarize_transcript_tool,
            description=(
                "Summarize the transcript by identifying key insights and main points. "
                "Input should be a JSON string with 'transcript' and 'speaker_mapping'."
            ),
        ),
        Tool(
            name="Extract Hard Facts",
            func=extract_hard_facts_tool,
            description=(
                "Extract hard, factual information mentioned in the transcript. "
                "Input should be the transcript as a single string."
            ),
        ),
        Tool(
            name="Extract Financial Advice",
            func=extract_financial_advice_tool,
            description=(
                "Extract financial advice or actionable recommendations from the transcript. "
                "Input should be the transcript as a single string."
            ),
        ),
        Tool(
            name="Extract Financial Goals",
            func=extract_financial_goals_tool,
            description=(
                "Extract financial goals from the client or financial adviser from the transcript. "
                "Input should be the transcript as a single string."
            ),
        ),
    ]

# Define your tool functions with a single input argument

def summarize_transcript_tool(input_json: str) -> str:
    """
    Summarizes the transcript with the given speaker mapping.
    Expects a JSON string with 'transcript' and 'speaker_mapping'.
    """
    try:
        input_data = json.loads(input_json)
        transcript = input_data['transcript']
        speaker_mapping = input_data['speaker_mapping']
        summary = summarize_transcript_with_mapping(transcript=transcript, speaker_mapping=speaker_mapping)
        return summary
    except (json.JSONDecodeError, KeyError) as e:
        print(f"Invalid input for summarize_transcript_tool: {e}")
        return "Error: Invalid input format for summarizing transcript."

def extract_hard_facts_tool(transcript: str) -> str:
    """
    Extracts hard facts from the transcript.
    """
    try:
        facts = extract_hard_facts_with_llm(transcript)
        return facts
    except Exception as e:
        print(f"Error in extract_hard_facts_tool: {e}")
        return "Error: Failed to extract hard facts."


def extract_financial_advice_tool(transcript: str) -> str:
    """
    Extracts financial advice from the transcript.
    """
    try:
        advice = extract_financial_advice_with_llm(transcript)
        return advice
    except Exception as e:
        print(f"Error in extract_financial_advice_tool: {e}")
        return "Error: Failed to extract financial advice."


def extract_financial_goals_with_llm(transcript: str) -> list:
    """Extracts financial goals from a transcript."""
    prompt = f"""You are an AI assistant tasked with extracting **financial goals** from a transcript.

    Financial goals are:
    - Stated objectives or aspirations mentioned by the **client**.
    - Can include saving, investing, retirement planning, debt reduction, homeownership, etc.
    - Must be **clearly stated by the client**, not inferred or assumed.
    - **Do not include** financial advice given by the planner.

    Your output must:
    - Be a **list of financial goals** extracted from the transcript.
    - Return "None" if no goals are mentioned.

    ### **Transcript:**
    {transcript}

    ### **Extracted Financial Goals:**
    - 
    """
    response_msg, _ = make_llm_query(prompt=prompt, model_name=SELECTED_MODEL, log_file=LLM_LOG_FILE_NAME)
    if response_msg.lower() == "none":
        return []
    return response_msg
  
  
def extract_financial_goals_tool(transcript: str) -> str:
    """
    Extracts financial advice from the transcript.
    """
    try:
        advice = extract_financial_goals_with_llm(transcript)
        return advice
    except Exception as e:
        print(f"Error in extract_financial_advice_tool: {e}")
        return "Error: Failed to extract financial advice."


# Initialize the agent once
def create_agent() -> Optional[AgentExecutor]:
    tools = get_tools()

    try:
        llm = LangChainOpenAI(temperature=0, max_tokens=1000)
        agent = initialize_agent(
            tools=tools,
            llm=llm,
            agent="zero-shot-react-description",
            verbose=True
        )
        print("Agent initialized successfully.")
        return agent
    except Exception as e:
        print(f"Failed to initialize the agent: {e}")
        return None
    
def process_audio_segment(chunk: List[Dict], speaker_mapping: Dict[str, str], agent) -> Dict[str, Optional[str]]:
    """
    Processes a single transcript chunk using the initialized agent to summarize, extract hard facts, and extract advice.

    Parameters:
    - chunk (list): A chunk of transcript segments.
    - speaker_mapping (dict): A dictionary mapping speakers to roles.
    - agent: The initialized Langchain agent.

    Returns:
    - dict: Processed results for the chunk, including summary, hard facts, and advice.
    """
    # Combine chunk into a single transcript text    
    pre_all_chunk_text = []
    for seg_text, speaker in zip(chunk['text'], chunk['speakers']):
        pre_all_chunk_text.append(f"{speaker['speaker']}: {seg_text}")
    
    all_chunk_text = "\n".join(pre_all_chunk_text)
    print(all_chunk_text)
        
    # Prepare JSON inputs for each tool
    summarize_input = json.dumps({
        "transcript": all_chunk_text,
        "speaker_mapping": speaker_mapping
    })
    
    hard_facts_input = json.dumps({
        "transcript": all_chunk_text
    })
    
    financial_advice_input = json.dumps({
        "transcript": all_chunk_text
    })
    financial_goals_input = json.dumps({
        "transcript": all_chunk_text
    })
    # Define the tool inputs
    tool_inputs = {
        "Summarize Transcript": summarize_input,
        "Extract Hard Facts": hard_facts_input,
        "Extract Financial Advice": financial_advice_input,
        "Extract Financial Goals": financial_goals_input,
    }

    results = {}
    # Iterate over each tool and execute its function with appropriate arguments
    for tool in agent.tools:
        try:
            input_for_tool = tool_inputs.get(tool.name, "")
            # Call other tools with just the transcript
            tool_result = tool.run(input_for_tool)
            
            # Store the result with a formatted key
            results[tool.name.lower().replace(" ", "_")] = tool_result
            print(f"Tool '{tool.name}' executed successfully.")
        except Exception as e:
            print(f"Error running tool '{tool.name}': {e}")
            results[tool.name.lower().replace(" ", "_")] = None
        
        # Try stop API rate limits
        time.sleep(0.1)

    return results



def generate_combined_summary(summaries: list) -> str:
    """
    Combines individual summaries into a cohesive final summary using OpenAI.

    Parameters:
    - summaries (list): List of individual chunk summaries.

    Returns:
    - str: A single cohesive summary.
    """
    summaries_text = "\n\n".join([s for s in summaries if s is not None])
    # prompt = f"""Combine the following summaries into a cohesive, comprehensive summary of a few sentences long. Text to summarize:

    # {summaries_text}
    
    # Ensure the final summary is concise, clear, and captures all key points from the individual summaries.
    # """
    prompt = f"""You are an AI assistant tasked with combining multiple summaries into one cohesive and concise summary. 

    Your goal is to:
    - Combine the following summaries into a single, comprehensive summary that captures all the key points.
    - Ensure the final summary is clear, concise, and only a few sentences long.
    - Remove all instances of 'Errors' or 'None' from the list.

    Individual summaries:
    {summaries_text}

    Final cohesive summary:
    """
    response_msg, cost_for_query = make_llm_query(prompt=prompt, model_name=SELECTED_MODEL, log_file=LLM_LOG_FILE_NAME)
    return response_msg, cost_for_query

def generate_combined_advice(advice_list: list) -> str:
    """
    Combines advice into a concise, deduplicated bullet-point list using OpenAI.

    Parameters:
    - advice_list (list): List of advice extracted from chunks.

    Returns:
    - str: Deduplicated and consolidated advice.
    """
    if len(advice_list) == 0:
        return [], 0

    advice_text = "\n".join([a for a in advice_list if a is not None])

    prompt = f"""You are an AI assistant tasked with deduplicating and consolidating financial advice into a clear, organized list.

    Your goal is to:
    - Combine the following advice points into a concise and actionable bullet-point list.
    - Deduplicate similar advice and ensure there is no overlap or redundancy.
    - Return the advice in the format: ['Advice 1', 'Advice 2', 'Advice 3'].
    - Remove all instances of 'Errors' or 'None' from the list.
    - If there is no clear advice, return 'None'.

    Advice points to consolidate:
    {advice_text}

    Final consolidated advice:
    """
    response_msg, cost_for_query = make_llm_query(prompt=prompt, model_name=SELECTED_MODEL, log_file=LLM_LOG_FILE_NAME)
    if response_msg.lower() == "none":
        return [], 0
    return response_msg, cost_for_query

def generate_combined_goals(goals_list: list) -> str:
    """
    Combines financial goals into a concise, deduplicated bullet-point list using OpenAI/Gemini.

    Parameters:
    - goals_list (list): List of financial goals extracted from chunks.

    Returns:
    - str: Deduplicated and consolidated financial goals.
    """
    if len(goals_list) == 0:
        return [], 0

    goals_text = "\n".join([g for g in goals_list if g is not None])

    prompt = f"""You are an AI assistant tasked with deduplicating and consolidating financial goals into a clear, organized list.

    Your goal is to:
    - Combine the following financial goals into a **concise, structured bullet-point list**.
    - Deduplicate similar goals and ensure there is **no overlap or redundancy**.
    - Return the goals in the format: **['Goal 1', 'Goal 2', 'Goal 3']**.
    - Remove all instances of **'Errors' or 'None'** from the list.
    - If there are no valid financial goals, return **'None'**.

    Financial goals to consolidate:
    {goals_text}

    Final consolidated financial goals:
    """
    
    response_msg, cost_for_query = make_llm_query(prompt=prompt, model_name=SELECTED_MODEL, log_file=LLM_LOG_FILE_NAME)
    
    if response_msg.lower() == "none":
        return [], 0
    
    return response_msg, cost_for_query

def combine_non_empty_sublists(lst):
    combined = [item for sublist in lst if sublist for item in sublist]
    return combined if combined else []

def generate_combined_hard_facts(hard_facts_list: list) -> str:
    """
    Combines hard facts into a concise, deduplicated bullet-point list using OpenAI.

    Parameters:
    - hard_facts_list (list): List of hard facts extracted from chunks.

    Returns:
    - str: Deduplicated and consolidated hard facts as a list.
    """
    all_hard_facts = combine_non_empty_sublists([s for s in hard_facts_list if s is not None])
    if len(all_hard_facts) == 0:
        return [], 0

    hard_facts_text = "\n".join(all_hard_facts)

    prompt = f"""You are an AI assistant tasked with consolidating and deduplicating a list of hard personal and financial facts. 
    A hard fact is a statement of objective truth, free from ambiguity, interpretation, or subjectivity. The hard facts should not be about the conversation that took place, 
    but rather about the financial information or personal information shared. 

    Your goal is to:
    - Extract only hard facts from the provided list, avoiding any overlap or redundancy.
    - Ensure each fact is unique and does not repeat information stated elsewhere in the list.
    - Exclude any subjective opinions, recommendations, or ambiguous statements.
    - Exclude any factual financial information that is not directly related to the client's personal or financial situation.
    - Present the final output in the format: ['Fact 1', 'Fact 2', 'Fact 3'].
    - Remove all instances of 'Errors' or 'None' from the list.
    - If no valid hard facts can be extracted, return 'None'.

    Hard facts to analyze and consolidate:
    {hard_facts_text}

    Final consolidated hard facts:
    """
    response_msg, cost_for_query = make_llm_query(prompt=prompt, model_name=SELECTED_MODEL, log_file=LLM_LOG_FILE_NAME)
    if response_msg.lower() == "none":
        return [], 0
    return response_msg, cost_for_query


def aggregate_chunk_results(chunk_results: list) -> dict:
    """
    Aggregates results from all processed chunks into a final summary, list of hard facts, 
    and optionally goals and advice.

    Parameters:
    - chunk_results (list): List of dictionaries containing processed results for each chunk.

    Returns:
    - dict: Aggregated results including final summary, hard facts, goals, and advice.
    """
    # Extract and prepare data for aggregation
    summaries = [chunk["summarize_transcript"] for chunk in chunk_results]
    all_advice = [chunk["extract_financial_advice"] for chunk in chunk_results if chunk["extract_financial_advice"]!= []]
    all_hard_facts_list = [chunk["extract_hard_facts"] for chunk in chunk_results]
    all_goal = [chunk["extract_financial_goals"] for chunk in chunk_results]
    print("all_hard_facts_list single entry:\n", all_hard_facts_list[0])
    
    # Generate consolidated outputs
    final_full_summary, _ = generate_combined_summary(summaries)
    final_combined_advice, _ = generate_combined_advice(all_advice)
    final_combined_hard_facts, _ = generate_combined_hard_facts(all_hard_facts_list)
    final_combined_goals, _ = generate_combined_goals(all_goal) 

    return {
        "final_full_summary": final_full_summary,
        "final_combined_advice": final_combined_advice,
        "all_hard_facts": final_combined_hard_facts,
        "all_financial_goals": final_combined_goals,
    }


import json
import tiktoken
from typing import List
from collections import Counter



def classify_roles_with_llm_V2(transcript: str, n_speakers: int, model_name: str) -> dict:
    """Classifies roles in the transcript, handling token limits and aggregating results."""
    token_limit = MODEL_TOKEN_LIMITS.get(model_name, 10000)
    transcript_tokens = count_tokens_func(transcript, model_name)
    print(f"Total Tokens in transcription: {transcript_tokens}")
    
    if transcript_tokens > token_limit:
        print(f"(ISSUE) Transcript too large ({transcript_tokens} tokens). Just using first chunk...")
        chunks = split_text_into_chunks(transcript, model_name)
        selected_chunk = chunks[0]
        tmp_count = count_tokens_func(selected_chunk, model_name)
        print(f"(ISSUE) Total Tokens in sub-selected transcription: {tmp_count}")
    else:
        selected_chunk = transcript

    classification_output = classify_speakers_api_call(transcript=selected_chunk, n_speakers=n_speakers)
    return classification_output    
    
    
def generate_speaker_mapping(combined_speaker_segments: List[dict]) -> dict:
    """
    Generates a speaker mapping by classifying roles in the entire transcript.

    Parameters:
    - raw_transcript (dict): The raw transcription output from AWS Transcribe.

    Returns:
    - dict: A dictionary mapping speaker IDs to their roles (e.g., "financial planner", "client").
    """
    # Combine all segments into a single transcript text
    transcript_text = "\n".join(
        [f"{seg['speaker_label']}: {seg['transcript']}" for seg in combined_speaker_segments]
    )
    n_speakers = len(set(seg['speaker_label'] for seg in combined_speaker_segments))
    print(f"n_speakers: {n_speakers}")

    # Classify roles using OpenAI
    # Ensure we select the biggest piece of text which can be used with the model
    classification_output = classify_roles_with_llm_V2(
        transcript=transcript_text, n_speakers=n_speakers,
        model_name=SELECTED_MODEL
    )
    print("classification_output", classification_output)
    return classification_output

    # Parse the classification output to generate the speaker mapping
    # speaker_mapping = parse_speaker_mapping(classification_output, combined_speaker_segments)
    # return speaker_mapping


def create_transcript_output_docx(combined_speaker_segments: List[dict], speaker_mapping: dict, meeting_name: str, fp_name="Morgan", client_name: str = "Client A"):
    """
    Create a formatted DOCX file from a AWS Transcribe output.

    Args:
        combined_speaker_segments (dict): List of Speaker Segments with speaker and transcript as dict keys
        speaker_mapping (dict): Mapped dictionary of speaker indices to speaker roles
        client_name (str): Name of the client.
        meeting_name (str): Name of the meeting.
        fp_name (str): Name of the FP (default is "Morgan").
    """
    # Initialize document
    document = Document()

    # Add metadata section
    document.add_heading(f"Client: {client_name}", level=1)
    document.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    document.add_paragraph(f"Meeting Name: {meeting_name}")
    document.add_paragraph(f"Financial Planner: {fp_name}")
    document.add_paragraph("_" * 50)
    document.add_paragraph("Speaker Mapping:")
    for speaker_label, speaker_name in speaker_mapping.items():
        document.add_paragraph(f"{chr(ord('A') + list(speaker_mapping.keys()).index(speaker_label))}: {speaker_name}")

    document.add_paragraph("\n")
    document.add_paragraph("Transcript")
    document.add_paragraph("_" * 50)

    # Add combined transcript to the document
    for segment in combined_speaker_segments:
        speaker_label = segment.get('speaker_label')
        transcript = segment.get('transcript')
        speaker_name = chr(ord('A') + list(speaker_mapping.keys()).index(speaker_label))  # Map spk_0, spk_1, etc. to A, B, C...
        document.add_paragraph(f"{speaker_name}: {transcript}")

    # Save the document
    filename = os.path.join(
        transcript_out_dir,
        f"transcript_{meeting_name}_{datetime.now().strftime('%d-%m-%y')}.docx"
    )
    document.save(filename)
    print(f"Transcript saved as {filename}")
    

def create_combined_speaker_segments(raw_transcript: dict) -> List[dict]:
    # Combine consecutive segments by the same speaker
    combined_segments = []
    previous_speaker = None
    current_transcript = ""

    for segment in raw_transcript.get('audio_segments', {}):
        speaker_label = segment.get('speaker_label')
        transcript = segment.get('transcript')

        if speaker_label == previous_speaker:
            current_transcript += f" {transcript}"  # Append to current transcript
        else:
            if previous_speaker is not None:
                combined_segments.append({"speaker_label": previous_speaker, "transcript": current_transcript})
            previous_speaker = speaker_label
            current_transcript = transcript

    # Add the last segment
    if previous_speaker is not None:
        combined_segments.append({"speaker_label": previous_speaker, "transcript": current_transcript})
        
    return combined_segments



def merge_transcript_speaker_breaks(raw_transcripts: dict) -> dict:
    """
    Merges consecutive segments in the AWS Transcribe output where the same speaker is speaking,
    updating timestamps and preserving the format.
    
    :param raw_transcripts: Dictionary containing AWS Transcribe raw output.
    :return: Updated dictionary with merged speaker segments.
    """
    merged_audio_segments = []
    current_segment = None
    current_items = []
    items = raw_transcripts["items"]

    # Ensure all items have an end_time by inferring from the next item's start_time
    for i in range(len(items) - 1):
        if "end_time" not in items[i]:
            items[i]["end_time"] = items[i + 1]["start_time"]
    if "end_time" not in items[-1]:  # Handle the last item
        items[-1]["end_time"] = items[-1].get("start_time")  # Keep it as start_time if nothing else

    for segment in raw_transcripts["audio_segments"]:
        # Start a new segment if it's the first or speaker changes
        if not current_segment or segment.get("speaker_label") != current_segment.get("speaker_label"):
            # Save the completed segment
            if current_segment:
                current_segment["start_time"] = current_items[0]["start_time"]
                current_segment["end_time"] = current_items[-1]["end_time"]
                current_segment["transcript"] = " ".join(
                    [item["alternatives"][0]["content"] for item in current_items]
                )
                current_segment["items"] = [item["id"] for item in current_items]
                merged_audio_segments.append(current_segment)

            # Start a new segment
            current_segment = {
                "id": len(merged_audio_segments),
                "speaker_label": segment.get("speaker_label"),
                "start_time": None,
                "end_time": None,
                "transcript": "",
                "items": []
            }
            current_items = []

        # Collect items for the current segment
        for item_id in segment["items"]:
            current_items.append(items[item_id])

    # Add the final segment
    if current_segment:
        current_segment["start_time"] = current_items[0]["start_time"]
        current_segment["end_time"] = current_items[-1]["end_time"]
        current_segment["transcript"] = " ".join(
            [item["alternatives"][0]["content"] for item in current_items]
        )
        current_segment["items"] = [item["id"] for item in current_items]
        merged_audio_segments.append(current_segment)

    # Update the raw_transcripts dictionary
    raw_transcripts["audio_segments"] = merged_audio_segments

    # Concatenate transcripts into logical sections for the 'transcripts' key
    concatenated_transcripts = []
    current_transcript = ""

    for segment in merged_audio_segments:
        # Append the current segment's transcript
        if current_transcript:
            current_transcript += " " + segment["transcript"]
        else:
            current_transcript = segment["transcript"]

        # Logic to end a concatenation group (based on speaker changes or other conditions)
        if "speaker_label" in segment:
            concatenated_transcripts.append(current_transcript)
            current_transcript = ""

    # Add the final group if any
    if current_transcript:
        concatenated_transcripts.append(current_transcript)

    # Update the transcripts key
    raw_transcripts["transcripts"] = [{"transcript": t} for t in concatenated_transcripts]
    return raw_transcripts



def extract_confidence_statistics(raw_transcription: dict):
    """
    Extracts confidence scores, calculates averages and standard deviations for:
    - 5-minute intervals
    - Total transcription confidence
    - Confidence by speaker

    Parameters:
    - raw_transcription (dict): Raw AWS Transcribe data.

    Returns:
    - dict: Statistics with confidence scores in 5-minute intervals, overall, and per speaker.
    """
    # Extract items and speaker information
    items = raw_transcription["items"]
    speaker_segments = raw_transcription["speaker_labels"]["segments"]

    # Collect confidence scores with timestamps
    confidence_scores = []
    for item in items:
        if "alternatives" in item and "confidence" in item["alternatives"][0]:
            start_time = float(item["start_time"]) if "start_time" in item else None
            end_time = float(item["end_time"]) if "end_time" in item else None
            confidence = float(item["alternatives"][0]["confidence"])
            confidence_scores.append({"start_time": start_time, "end_time": end_time, "confidence": confidence})

    # Calculate overall confidence statistics
    confidences = [score["confidence"] for score in confidence_scores]
    total_avg_confidence = np.mean(confidences)
    total_std_confidence = np.std(confidences)

    # 5-minute interval confidence statistics
    interval_confidences = []
    interval_stats = {}
    current_interval = 0
    for score in confidence_scores:
        interval_start = current_interval * 300
        interval_end = (current_interval + 1) * 300

        if score["start_time"] and interval_start <= score["start_time"] < interval_end:
            interval_confidences.append(score["confidence"])
        elif score["start_time"] and score["start_time"] >= interval_end:
            # Process the completed interval
            if interval_confidences:
                avg_confidence = np.mean(interval_confidences)
                std_confidence = np.std(interval_confidences)
                interval_stats[f"Interval {current_interval}"] = {"avg": avg_confidence, "std": std_confidence}
                interval_confidences = []
            # Move to the next interval
            current_interval += 1
            if current_interval * 300 <= score["start_time"] < (current_interval + 1) * 300:
                interval_confidences.append(score["confidence"])

    # Add the last interval if not empty
    if interval_confidences:
        avg_confidence = np.mean(interval_confidences)
        std_confidence = np.std(interval_confidences)
        interval_stats[f"Interval {current_interval}"] = {"avg": avg_confidence, "std": std_confidence}

    # Speaker-specific confidence statistics
    speaker_confidences = {}
    for segment in speaker_segments:
        speaker_label = segment["speaker_label"]
        segment_start = float(segment["start_time"])
        segment_end = float(segment["end_time"])
        segment_scores = [
            score["confidence"]
            for score in confidence_scores
            if score["start_time"] and segment_start <= score["start_time"] < segment_end
        ]
        if speaker_label not in speaker_confidences:
            speaker_confidences[speaker_label] = []
        speaker_confidences[speaker_label].extend(segment_scores)

    speaker_stats = {
        speaker: {
            "avg": np.mean(scores),
            "std": np.std(scores),
            "count": len(scores),
        }
        for speaker, scores in speaker_confidences.items()
        if scores
    }

    # Print results
    print("Total Transcription Confidence:")
    print(f"Average: {total_avg_confidence:.2f}, Standard Deviation: {total_std_confidence:.2f}\n")
    print("5-Minute Interval Confidence:")
    for interval, stats in interval_stats.items():
        print(f"{interval}: Avg={stats['avg']:.2f}, Std={stats['std']:.2f}")
    print("\nSpeaker Confidence:")
    for speaker, stats in speaker_stats.items():
        print(f"{speaker}: Avg={stats['avg']:.2f}, Std={stats['std']:.2f}, Count={stats['count']}")

    return {
        "total_confidence": {"avg": total_avg_confidence, "std": total_std_confidence},
        "interval_confidence": interval_stats,
        "speaker_confidence": speaker_stats,
    }


if __name__ == "__main__":
    # Step 1: Initialize clients and managers
    boto3_client, s3_manager, transcription_manager = initialize_clients(
        aws_access_key_id=AWS_ACCESS_KEY_ID,
        aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
    )
    print(f"Initialized Client Connections Successfully")
    if not DRY_RUN_SYSTEM:
        # Step 2: Upload audio
        # local_file_path = "uploads/api-test1/client_conversation.wav"
        # local_file_path = "uploads/FP Chat Morgan - 19 Jan 2025/Partial-1-18min-Audio-Clip-Ryan-Mccarlie.wav"
        local_file_path = "uploads/FP Chat Morgan - 19 Jan 2025/Full-Audio-Clip-Ryan-Mccarlie.wav"
        s3_file_uri, s3_output_folder = upload_audio_to_s3(s3_manager=s3_manager, local_file_path=local_file_path)
        print(f"Uploaded S3 file uri: {s3_file_uri}. S3 Output Path: {s3_output_folder}")
        
        try:
            # Step 3: Start transcription job
            job_name = start_aws_transcription_job(
                transcription_manager=transcription_manager,
                input_s3_uri=s3_file_uri,
                output_s3_dir_prefix=s3_output_folder,
            )
            print(f"Job name: {job_name}")
            
            # Step 4: Wait for transcription to complete
            response = transcription_manager.wait_for_transcription_completion(job_name)
        
        except Exception as e:
            print(f"Error starting transcription job: {e}")
            raise e

    else:
        job_name = "transcription_job_20250218_215916"
        s3_output_folder = "other/audio-files/job_2158_18-02-25/output/"

    # Step 5: Fetch transcription output
    raw_transcripts = transcription_manager.fetch_transcription_output(job_name, s3_output_folder=s3_output_folder)
    print(f"Successfully fetched raw transcriptions")

    confidence_stats = extract_confidence_statistics(raw_transcription=raw_transcripts)
    print("confidence_stats", confidence_stats)

    merged_transcripts = merge_transcript_speaker_breaks(raw_transcripts=raw_transcripts)
    combined_speaker_segments = merged_transcripts["audio_segments"]
    
    # Generate Speaker Mapping like so
    speaker_mapping = generate_speaker_mapping(combined_speaker_segments=combined_speaker_segments)
    # speaker_mapping = {'spk_0': 'financial planner', 'spk_1': 'client'}
    print("speaker_mapping:\n", speaker_mapping)    
        
    meeting_name = "Intro FP Chat with Ryan"
    client_name = "Ryan"
    create_transcript_output_docx(
        combined_speaker_segments=combined_speaker_segments, 
        speaker_mapping=speaker_mapping,
        meeting_name=meeting_name, client_name=client_name, fp_name=DEFAULT_FP_NAME
    )

    # Step 2: Chunk the transcript into manageable pieces
    transcription_chunks = partition_transcription_into_chunks(speaker_segment_transcriptions=combined_speaker_segments, chunk_duration_s=300)
    print(f"Number of transcription_chunks produced: {len(transcription_chunks)}")
    print(f"Size of first chunk: {len(transcription_chunks[0])}")
    print(f"First chunk: {transcription_chunks[0]}")    
    
    # Initialize tools and agent
    agent = create_agent()

    if agent is None:
        raise RuntimeError("Agent initialization failed. Cannot process audio segments.")
    
    # Step 3: Process each chunk    
    processed_audio_results = []
    for i, audio_segments in enumerate(transcription_chunks):
        processed_results = process_audio_segment(
            agent=agent,chunk=audio_segments,speaker_mapping=speaker_mapping
        )
        processed_audio_results.append(processed_results)
        time.sleep(0.1)

    # Step 4: Aggregate the results from all chunks
    final_results = aggregate_chunk_results(processed_audio_results)

    # # Output final result
    print(json.dumps(final_results, indent=4))
    
    # Get full query stats
    stats = cost_tracker.get_query_stats()
    
    # Log the final results
    log_final_results(final_results, stats, log_final_file=FULL_APP_LOG_FILE_NAME)

    print("Query Statistics:")
    for key, value in stats.items():
        print(f"{key}: {value}")
